from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ======================================================
# 1. CÁC HÀM HỖ TRỢ XỬ LÝ TEXT & FORMATTING
# ======================================================

def apply_markdown_formatting(paragraph, text):
    """
    Xử lý in đậm (**text**) và in nghiêng (*text*) cho một đoạn văn (paragraph).
    Hàm này dùng chung cho cả nội dung bài viết và các ô trong bảng.
    """
    if not text: return

    # Tách phần In đậm (**) trước
    parts = text.split('**')
    for i, part in enumerate(parts):
        # Nếu i lẻ -> Nằm giữa cặp ** -> Xử lý In Đậm
        if i % 2 != 0: 
            run = paragraph.add_run(part)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Nếu i chẵn -> Văn bản thường hoặc chứa In Nghiêng (*)
        else:
            # Tách tiếp phần In nghiêng (*)
            sub_parts = part.split('*')
            for j, sub_part in enumerate(sub_parts):
                run = paragraph.add_run(sub_part)
                # Nếu j lẻ -> Nằm giữa cặp * -> Xử lý In Nghiêng
                if j % 2 != 0:
                    run.italic = True

def clean_cell_text(text):
    """Làm sạch văn bản trong ô"""
    return text.strip()

def is_table_row(text):
    """Kiểm tra xem dòng này có phải là bảng không (chứa ký tự |)"""
    return '|' in text

def is_separator_line(row_data):
    """Kiểm tra dòng phân cách Markdown (---|---|---)"""
    clean_row = [str(cell).strip().replace('-', '') for cell in row_data]
    return all(c == '' for c in clean_row)

# ======================================================
# 2. CÁC HÀM XỬ LÝ LOGIC BẢNG (MERGE & DRAW)
# ======================================================

def parse_table_block(block_lines):
    """
    Phân tích khối văn bản (list các dòng string) thành dữ liệu bảng (List of Lists)
    """
    table_data = []
    
    for line in block_lines:
        clean_line = line.strip()
        if not clean_line: continue

        if '|' in clean_line:
            # Tách cột
            cells = [c for c in clean_line.split('|')]
            
            # Markdown thường có | ở đầu và cuối, nếu split ra sẽ có phần tử rỗng ở 2 đầu -> remove
            if clean_line.startswith('|') and len(cells) > 0: cells.pop(0)
            if clean_line.endswith('|') and len(cells) > 0: cells.pop(-1)
            
            # Làm sạch khoảng trắng từng ô
            cells = [clean_cell_text(c) for c in cells]

            # Bỏ qua dòng kẻ ngang phân cách (---|---)
            if is_separator_line(cells):
                continue
                
            table_data.append(cells)
        else:
            # Xử lý rớt dòng (nối vào ô cuối của dòng trước nếu dòng này không có |)
            if table_data:
                table_data[-1][-1] += " " + clean_line

    return table_data

def process_fill_down(data):
    """Tự động điền dữ liệu cho cột đầu tiên nếu trống (Auto Fill Down)"""
    if not data: return data
    for i in range(1, len(data)):
        current_row = data[i]
        prev_row = data[i-1]
        if len(current_row) > 0:
            first_cell = current_row[0]
            if not first_cell and len(prev_row) > 0:
                data[i][0] = prev_row[0]
    return data

def draw_word_table(doc, data):
    """Vẽ bảng Word và áp dụng Format Markdown cho từng ô"""
    if not data: return

    data = process_fill_down(data)

    max_cols = 0
    for row in data:
        if len(row) > max_cols: max_cols = len(row)
    if max_cols == 0: return

    table = doc.add_table(rows=len(data), cols=max_cols)
    table.style = 'Table Grid'

    for i, row_data in enumerate(data):
        row = table.rows[i]
        safe_row_data = row_data + [""] * (max_cols - len(row_data))
        
        for j, cell_text in enumerate(safe_row_data):
            if j < len(row.cells):
                # Lấy paragraph mặc định trong ô
                p = row.cells[j].paragraphs[0]
                # Áp dụng hàm format (để xử lý **bold**, *italic*) thay vì gán .text thuần
                apply_markdown_formatting(p, str(cell_text))
    
    doc.add_paragraph()

# ======================================================
# 3. HÀM CHÍNH: CREATE DOCUMENT
# ======================================================

def create_final_docx(path, char_name, intro_text, infobox_raw, body_text, original_url):
    """Tạo file Word cuối cùng"""
    doc = Document()
    
    # --- 1. TITLE ---
    main_title = doc.add_heading(f"{char_name} | Wiki nhân vật Game of Thrones", level=0)
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\n")

    # --- 2. INTRO ---
    if intro_text:
        p = doc.add_paragraph() 
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        apply_markdown_formatting(p, intro_text)
    
    doc.add_paragraph("\n")

    # --- 3. INFOBOX TABLE ---
    if infobox_raw:
        doc.add_heading(f'Tổng quan về {char_name}', level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.rows[0].cells[0].text = 'Thông tin'
        table.rows[0].cells[1].text = 'Chi Tiết'
        
        for line in infobox_raw.split('\n'):
            if '|' in line:
                key, val = line.split('|', 1)
                if key.strip() and val.strip():
                    row = table.add_row()
                    row.cells[0].text = f"✅ {key.strip()}"
                    row.cells[1].text = f"⭐ {val.strip()}"
        doc.add_paragraph("\n")

    # --- 4. BODY (AUTO TABLE DETECT) ---
    doc.add_heading('Nội dung chi tiết', level=1)
    
    lines = body_text.split('\n')
    table_buffer = [] 
    
    for i, line in enumerate(lines):
        line = line.strip()
        
        # [A] LOGIC BẢNG
        if is_table_row(line):
            table_buffer.append(line)
            
            is_end_of_table = False
            if i == len(lines) - 1:
                is_end_of_table = True
            elif not is_table_row(lines[i+1].strip()):
                is_end_of_table = True
            
            if is_end_of_table:
                data_matrix = parse_table_block(table_buffer)
                draw_word_table(doc, data_matrix)
                table_buffer = []
            continue
        
        # [B] LOGIC TEXT THƯỜNG
        if not line: continue
        
        if line.startswith('### '): doc.add_heading(line[4:], level=3)
        elif line.startswith('## '): doc.add_heading(line[3:], level=2)
        elif line.startswith('# '): doc.add_heading(line[2:], level=1)
        else:
            p = doc.add_paragraph()
            apply_markdown_formatting(p, line)

    # --- 5. FOOTER ---
    doc.add_paragraph("\n")
    doc.add_heading('Thư Viện Ảnh', level=2)
    doc.add_paragraph('[gallery link="file" columns="2" size="medium" ids="..."]')
    doc.add_paragraph("---")
    
    footer = doc.add_paragraph()
    footer.add_run(f'Các bạn đang theo dõi bài viết “{char_name} | Wiki nhân vật Game of Thrones”...')
    footer.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    ref = doc.add_paragraph()
    ref.add_run("Nguồn tham khảo:").bold = True
    ref.add_run(f"\n{char_name} - {original_url}")

    doc.save(path)