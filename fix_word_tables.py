import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
import os
import re

def is_table_row(text):
    """Kiểm tra xem dòng này có phải là bảng không"""
    return '|' in text

def is_separator_line(row_data):
    """Kiểm tra dòng phân cách (---|---|---)"""
    clean_row = [cell.strip().replace('-', '') for cell in row_data]
    return all(c == '' for c in clean_row)

def clean_cell_text(text):
    """Làm sạch văn bản trong ô"""
    return text.strip()

def parse_table_block(block_lines):
    """
    Phân tích khối văn bản thành dữ liệu bảng (List of Lists)
    """
    table_data = []
    
    for line in block_lines:
        clean_line = line.strip()
        if not clean_line: continue

        if '|' in clean_line:
            # Tách cột
            cells = [c for c in clean_line.split('|')]
            
            # Xử lý 2 đầu mút nếu có dấu |
            if clean_line.startswith('|') and len(cells) > 0: cells.pop(0)
            if clean_line.endswith('|') and len(cells) > 0: cells.pop(-1)
            
            # Làm sạch khoảng trắng từng ô
            cells = [clean_cell_text(c) for c in cells]

            # Bỏ qua dòng kẻ ngang
            if is_separator_line(cells):
                continue
                
            table_data.append(cells)
        else:
            # Xử lý rớt dòng (nối vào ô cuối của dòng trước)
            if table_data:
                table_data[-1][-1] += " " + clean_line

    return table_data

def process_fill_down(data):
    """
    Tự động điền dữ liệu cho cột đầu tiên nếu nó bị trống.
    Ví dụ: 
    Row 1: | Giải A | 2011 | ...
    Row 2: |        | 2012 | ...  -> Sẽ tự điền 'Giải A' vào ô trống
    """
    if not data: return data

    # Duyệt từ dòng thứ 2 trở đi
    for i in range(1, len(data)):
        current_row = data[i]
        prev_row = data[i-1]
        
        # Kiểm tra nếu dòng hiện tại có dữ liệu nhưng cột đầu tiên rỗng
        if len(current_row) > 0:
            first_cell = current_row[0]
            
            # Nếu cột 1 rỗng VÀ dòng trước đó có dữ liệu
            if not first_cell and len(prev_row) > 0:
                # Copy dữ liệu từ dòng trên xuống
                data[i][0] = prev_row[0]
    
    return data

def create_table_in_doc(doc, data):
    """Tạo bảng Word từ dữ liệu"""
    if not data: return

    # Bước tiền xử lý: Fill down dữ liệu trống ở cột 1
    data = process_fill_down(data)

    # Tính số cột tối đa
    max_cols = max(len(row) for row in data)
    
    # Tạo bảng
    table = doc.add_table(rows=len(data), cols=max_cols)
    table.style = 'Table Grid'

    for i, row_data in enumerate(data):
        row = table.rows[i]
        
        # Nếu dòng hiện tại ít cột hơn max_cols (do lỗi format), tự động thêm ô trống vào list data ảo để không lỗi index
        safe_row_data = row_data + [""] * (max_cols - len(row_data))
        
        for j, cell_text in enumerate(safe_row_data):
            if j < len(row.cells):
                row.cells[j].text = str(cell_text)
                
    # Thêm một dòng trống sau bảng để tách biệt
    doc.add_paragraph()

def process_file(filepath):
    try:
        doc = Document(filepath)
        new_doc = Document()
        
        buffer = [] 
        in_table_mode = False
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            # Nếu gặp dòng bắt đầu bằng |, hoặc đang trong bảng mà gặp text (kể cả rỗng để check fill down)
            if is_table_row(text):
                in_table_mode = True
                buffer.append(text)
            elif in_table_mode and text:
                # Dòng text thường nằm ngay sau bảng -> có thể là nội dung bị rớt dòng
                buffer.append(text)
            else:
                # Kết thúc bảng
                if buffer:
                    data = parse_table_block(buffer)
                    create_table_in_doc(new_doc, data)
                    buffer = []
                    in_table_mode = False
                
                # Ghi nội dung không phải bảng
                if text:
                    new_doc.add_paragraph(para.text)

        # Xử lý buffer cuối cùng
        if buffer:
            data = parse_table_block(buffer)
            create_table_in_doc(new_doc, data)

        # Lưu file
        dir_name = os.path.dirname(filepath)
        base_name = os.path.basename(filepath)
        name, ext = os.path.splitext(base_name)
        new_filename = f"{name}_fixed_v2{ext}"
        save_path = os.path.join(dir_name, new_filename)
        
        new_doc.save(save_path)
        return True, save_path
        
    except Exception as e:
        return False, str(e)

def select_files():
    files = filedialog.askopenfilenames(
        title="Chọn file Word lỗi",
        filetypes=[("Word Files", "*.docx")]
    )
    if not files: return

    count = 0
    errors = []
    
    for f in files:
        success, msg = process_file(f)
        if success:
            count += 1
        else:
            errors.append(f"{os.path.basename(f)}: {msg}")
    
    res_msg = f"Đã xử lý xong {count}/{len(files)} file."
    if errors:
        res_msg += "\n\nLỗi:\n" + "\n".join(errors)
        
    messagebox.showinfo("Hoàn tất", res_msg)

if __name__ == "__main__":
    root = tk.Tk()
    root.title("Tool Fix Bảng Word (Auto Fill)")
    root.geometry("350x150")
    
    lbl = tk.Label(root, text="Sửa lỗi bảng Word & Tự động điền dòng trống", pady=20)
    lbl.pack()
    
    btn = tk.Button(root, text="Chọn File Word", command=select_files, height=2, bg="#2196F3", fg="white")
    btn.pack()
    
    root.mainloop()